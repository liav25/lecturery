import os
from datetime import datetime

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

SUMMARY_FILE_NAME = "refined_full_summary.txt"


class LectureSummary:
    def __init__(self, directory):
        self.directory = directory
        self.date = self._parse_date(directory)
        self.file_path = os.path.join(directory, "transcriptions", SUMMARY_FILE_NAME)
        self.content = self._read_content()

    def _parse_date(self, directory):
        _, year, month, day = directory.split("_")
        return datetime(int(year), int(month), int(day))

    def _read_content(self):
        if not os.path.exists(self.file_path):
            return ""
        with open(self.file_path, "r", encoding="utf-8") as file:
            return file.read()


class DocxFormatter:
    def __init__(self):
        self.doc = Document()
        self._create_main_title_style()

    def _create_main_title_style(self):
        styles = self.doc.styles
        style = styles.add_style("MainTitle", WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(18)
        style.font.bold = True

    def add_lecture_title(self, lecture_num, date):
        formatted_date = date.strftime("%Y_%m_%d")
        self.doc.add_paragraph(f"Lecture {lecture_num}, {formatted_date}", style="MainTitle")
        self.doc.add_paragraph()

    def add_content(self, content):
        for line in content.split("\n"):
            if line.startswith("#"):
                level = len(line.split()[0])
                title = line.strip("#").strip()
                self.doc.add_heading(title, level=level)
            elif line.strip().startswith("- "):
                self.doc.add_paragraph(line.strip("- ").strip(), style="List Bullet")
            elif line.strip().startswith("    - "):
                self.doc.add_paragraph(line.strip("    - ").strip(), style="List Bullet 2")
            elif line.strip():
                self.doc.add_paragraph(line.strip())
            else:
                self.doc.add_paragraph()

    def add_page_break(self):
        self.doc.add_page_break()

    def save(self, filename):
        self.doc.save(filename)


def get_lecture_directories(home_path):
    return [
        os.path.join(home_path, d)
        for d in os.listdir(home_path)
        if os.path.isdir(os.path.join(home_path, d)) and d.startswith("lecture_")
    ]


def process_lectures(lecture_dirs):
    lectures = [LectureSummary(dir) for dir in lecture_dirs]
    return sorted(lectures, key=lambda x: x.date)


def create_combined_document(lectures, output_file):
    formatter = DocxFormatter()
    for i, lecture in enumerate(lectures, 1):
        formatter.add_lecture_title(i, lecture.date)
        formatter.add_content(lecture.content)
        formatter.add_page_break()
    formatter.save(output_file)


def main(home_path, output_file):
    lecture_dirs = get_lecture_directories(home_path)
    lectures = process_lectures(lecture_dirs)
    create_combined_document(lectures, output_file)
    print(f"Combined document saved as {output_file}")


if __name__ == "__main__":
    home_path = "data"  # Change this to the path where your lecture folders are located
    output_file = "combined_lecture_summaries.docx"
    main(home_path, output_file)
